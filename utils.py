from docx import Document


abbreviations = ["St.", "Mr.", "Ms.", "M.", "Dr.", "Mme.", "L.", "B.", "C.", "D.", "F.", "G.", "H.", "J.", "K.", "L.",
                 "N.", "P.", "Q.", "R.", "S.", "T.", "X.", "Z.", "W.", "V."]


def is_abbreviation(word):
    return word in abbreviations


def process_document(input_file, output_file):
    input_doc = Document(input_file)
    out_doc = Document()
    table = out_doc.add_table(rows=0, cols=1)

    for paragraph in input_doc.paragraphs:
        sentences = paragraph.text.split('.')
        i = 0
        while len(sentences[i]):
            # добавляем точку в конце каждого предложения
            sentence = sentences[i]
            if len(sentence):
                sentence += '.'
                # если предложение заканчивается на аббревиатуру - приклеиваем к нему следующее предложение
                # и меняем индекс
                words = sentence.split(' ')
                if is_abbreviation(words[-1]):
                    sentence = sentence + ' ' + sentences[i+1]
                    sentences[i] = sentence
                    del sentences[i+1]

                else:
                    i += 1

        for sentence in sentences:
            if len(sentence) > 0:
                cell = table.add_row().cells[0]
                cell.text = sentence + '.'

    out_doc.save(output_file)
